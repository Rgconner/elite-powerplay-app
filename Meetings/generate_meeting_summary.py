#!/usr/bin/env python3
"""
Generate Word document with meeting summary and SWOT analysis
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p

def create_meeting_summary():
    """Create the Word document with meeting summary"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Technical Meeting Summary & Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Meeting Overview
    add_heading(doc, 'Meeting Overview', 1)
    doc.add_paragraph(
        'Discussion between team members (including Dom, Jason, Russ, and Bob) regarding AI strategy, '
        'data governance, and SAP integration challenges at their organization (appears to be Bungie/related entity).'
    )
    
    doc.add_paragraph()  # Spacing
    
    # Key Discussion Topics
    add_heading(doc, 'Key Discussion Topics', 1)
    
    # Topic 1
    add_heading(doc, '1. AI Platform Strategy ("Three Flavors")', 2)
    doc.add_paragraph('Organization exploring three approaches to using Anthropic/Claude:')
    doc.add_paragraph('As a processor', style='List Bullet')
    doc.add_paragraph('Direct API integration', style='List Bullet')
    doc.add_paragraph('Alternative implementation methods', style='List Bullet')
    doc.add_paragraph('Still developing a cohesive AI strategy')
    
    # Topic 2
    add_heading(doc, '2. SAP API Policy Challenge', 2)
    p = doc.add_paragraph()
    run = p.add_run('Critical Issue: ')
    run.bold = True
    p.add_run('SAP announced strict policy prohibiting:')
    
    doc.add_paragraph('Third-party AI platforms from accessing SAP data', style='List Bullet')
    doc.add_paragraph('RPA (Robotic Process Automation) tools', style='List Bullet')
    doc.add_paragraph('Custom workarounds, middleware, proxies, or impersonation techniques', style='List Bullet')
    doc.add_paragraph('Autonomous or generative AI systems interacting with SAP APIs', style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Exception: ')
    run.bold = True
    p.add_run('Only SAP-endorsed architectures will be permitted')
    
    p = doc.add_paragraph()
    run = p.add_run('Implication: ')
    run.bold = True
    p.add_run('SAP likely preparing to announce their own AI architecture solution (vendor lock-in strategy)')
    
    doc.add_paragraph(
        'Team dubbed this the "AI Hunger Games" - vendors creating data fiefdoms'
    )
    
    # Topic 3
    add_heading(doc, '3. Data Governance & Quality Focus', 2)
    p = doc.add_paragraph()
    run = p.add_run('Primary Initiative: ')
    run.bold = True
    p.add_run('Establishing data quality and data inventory as foundational requirements')
    
    p = doc.add_paragraph()
    run = p.add_run('Programs Being Developed:')
    run.bold = True
    
    doc.add_paragraph('Data literacy programs', style='List Bullet')
    doc.add_paragraph('Data quality controls', style='List Bullet')
    doc.add_paragraph('Data inventory tools', style='List Bullet')
    doc.add_paragraph('Data lineage tracking (source to data product)', style='List Bullet')
    doc.add_paragraph('Business term classification', style='List Bullet')
    doc.add_paragraph('Completeness assessments', style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Philosophy: ')
    run.bold = True
    p.add_run('"Generally Regarded As Safe" (GRAS) data concept - publishing safe datasets for development without requiring individual permissions')
    
    # Topic 4
    add_heading(doc, '4. AI Governance & Orchestration', 2)
    doc.add_paragraph('Need for governance framework before exposing data to AI agents', style='List Bullet')
    doc.add_paragraph('Discussion of RAG (Retrieval-Augmented Generation) and other AI approaches', style='List Bullet')
    doc.add_paragraph('Requirement for gatekeepers when accessing real production data', style='List Bullet')
    doc.add_paragraph('Balance between enabling innovation and maintaining control', style='List Bullet')
    
    doc.add_page_break()
    
    # Next Steps
    add_heading(doc, 'Next Steps', 1)
    
    add_heading(doc, 'Immediate Actions:', 2)
    doc.add_paragraph(
        'Schedule follow-up meeting - May 6th (afternoon, 30 minutes) between Dom and team (Russ confirmed attendance)',
        style='List Number'
    )
    doc.add_paragraph(
        'Data governance discussion - Overview of data governance capabilities and publishing strategies',
        style='List Number'
    )
    doc.add_paragraph(
        'Identify technical counterpart - Dom needs to determine appropriate tech-side contact for deeper technical discussions',
        style='List Number'
    )
    doc.add_paragraph(
        "Monitor SAP developments - Track SAP's forthcoming AI architecture announcement",
        style='List Number'
    )
    doc.add_paragraph(
        'Continue data inventory work - Ongoing implementation of data inventory tools',
        style='List Number'
    )
    doc.add_paragraph(
        'Potential in-person meeting - Following week in Chesterfield if Bob is available',
        style='List Number'
    )
    
    add_heading(doc, 'Strategic Priorities:', 2)
    doc.add_paragraph('Finalize AI strategy approach', style='List Bullet')
    doc.add_paragraph('Establish data quality controls before AI implementation', style='List Bullet')
    doc.add_paragraph('Develop data literacy programs across organization', style='List Bullet')
    doc.add_paragraph('Create orchestration framework for AI tools', style='List Bullet')
    doc.add_paragraph('Navigate SAP restrictions and evaluate alternatives', style='List Bullet')
    
    doc.add_page_break()
    
    # SWOT Analysis
    add_heading(doc, 'SWOT Analysis', 1)
    
    # Strengths
    add_heading(doc, 'Strengths', 2)
    doc.add_paragraph(
        'Proactive approach to data governance and quality',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Recognition that data quality is foundational to AI success',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Team alignment on importance of data inventory and governance',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Vendor relationships with partners who understand both data governance and AI orchestration',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Early awareness of SAP restrictions before full implementation',
        style='List Bullet'
    )
    
    # Weaknesses
    add_heading(doc, 'Weaknesses', 2)
    doc.add_paragraph(
        'Lack of cohesive AI strategy - still exploring "three flavors" without clear direction',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Organizational structure gaps - unclear technical counterparts for key discussions',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Data maturity - described as "lone voice in the wilderness" initially on data quality',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Reactive posture to vendor policies (SAP) rather than proactive planning',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Limited current data governance - programs still in planning/early stages',
        style='List Bullet'
    )
    
    # Opportunities
    add_heading(doc, 'Opportunities', 2)
    doc.add_paragraph(
        'First-mover advantage in establishing proper data governance before competitors',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Partner expertise available for both data governance and AI orchestration',
        style='List Bullet'
    )
    doc.add_paragraph(
        'GRAS data concept could accelerate safe AI experimentation',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Cross-functional collaboration emerging between business and technical teams',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Market timing - establishing governance while AI adoption is still early',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Alternative to SAP lock-in - opportunity to evaluate other platforms before SAP forces decision',
        style='List Bullet'
    )
    
    # Threats
    add_heading(doc, 'Threats', 2)
    doc.add_paragraph(
        "Vendor lock-in - SAP's restrictive policies forcing architectural decisions",
        style='List Bullet'
    )
    doc.add_paragraph(
        'Competitive pressure - "AI Hunger Games" with vendors creating walled gardens',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Implementation delays - data governance work could slow AI adoption',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Resource constraints - multiple initiatives (data quality, inventory, literacy, governance) competing for attention',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Regulatory uncertainty - evolving AI compliance landscape',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Skills gap - need for data literacy programs suggests current capability gaps',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Integration complexity - orchestrating multiple AI tools and data sources',
        style='List Bullet'
    )
    
    # Save document
    filename = 'Meeting_Summary_and_SWOT_Analysis.docx'
    doc.save(filename)
    print(f"Document created successfully: {filename}")
    return filename

if __name__ == "__main__":
    create_meeting_summary()

# Made with Bob
